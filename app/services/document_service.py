"""
Document management service for the HR AI Assistant.

This service handles document processing, file management, and document requests.
"""

import os
import secrets
import mimetypes
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from pathlib import Path
from sqlalchemy.orm import Session

from app.models.document import Document, DocumentRequest, DocumentStatus, RequestStatus
from app.models.employee import Employee
from app.services.rag_service import rag_service
from app.utils.logger import get_logger

logger = get_logger(__name__)

class DocumentService:
    """Document management service"""
    
    def __init__(self):
        self.upload_directory = Path(os.getenv("UPLOAD_DIRECTORY", "./uploads"))
        self.max_file_size = int(os.getenv("MAX_FILE_SIZE", "10485760")) 
        self.allowed_extensions = os.getenv("ALLOWED_EXTENSIONS", ".pdf,.docx,.txt,.xlsx").split(",")
        
        # Ensure upload directory exists
        self.upload_directory.mkdir(parents=True, exist_ok=True)
    
    def process_uploaded_file(self, db: Session, file: BinaryIO, filename: str, 
                            author: Employee, **metadata) -> Tuple[bool, Optional[Document], List[str]]:
        """
        Process uploaded file and create document record
        
        Args:
            db: Database session
            file: File object
            filename: Original filename
            author: User uploading the file
            **metadata: Additional document metadata
            
        Returns:
            Tuple[bool, Optional[Document], List[str]]: (success, document, errors)
        """
        try:
            # Validate file
            validation_result = self._validate_file(file, filename)
            if not validation_result[0]:
                return False, None, validation_result[1]
            
            # Generate unique filename
            file_extension = Path(filename).suffix.lower()
            unique_filename = f"{secrets.token_hex(16)}{file_extension}"
            file_path = self.upload_directory / unique_filename
            
            # Save file
            file.seek(0)
            with open(file_path, "wb") as f:
                f.write(file.read())
            
            file_size = file_path.stat().st_size
            mime_type, _ = mimetypes.guess_type(filename)
            
            # Extract text content
            content_text = self._extract_text_content(file_path, file_extension)
            
            # Create document record
            document = Document(
                title=metadata.get("title", Path(filename).stem),
                description=metadata.get("description"),
                document_type=metadata.get("document_type"),
                file_path=str(file_path),
                file_name=filename,
                file_size=file_size,
                file_extension=file_extension,
                mime_type=mime_type,
                content_text=content_text,
                keywords=metadata.get("keywords"),
                tags=metadata.get("tags"),
                version=metadata.get("version", "1.0"),
                language=metadata.get("language", "en"),
                access_level=metadata.get("access_level"),
                author_id=author.id,
                status=DocumentStatus.DRAFT
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Index document for search if content is available
            if content_text and document.is_searchable:
                try:
                    rag_service.index_document(document, content_text)
                    document.opensearch_indexed = True
                    db.commit()
                except Exception as e:
                    logger.warning(f"Failed to index document {document.id}: {e}")
            
            logger.info(f"Document uploaded successfully: {document.id}")
            return True, document, []
            
        except Exception as e:
            logger.error(f"Error processing uploaded file: {e}")
            return False, None, [f"Error processing file: {str(e)}"]
    
    def _validate_file(self, file: BinaryIO, filename: str) -> Tuple[bool, List[str]]:
        """Validate uploaded file"""
        errors = []
        
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.allowed_extensions:
            errors.append(f"File type {file_extension} not allowed. Allowed types: {', '.join(self.allowed_extensions)}")
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > self.max_file_size:
            errors.append(f"File size ({file_size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)")
        
        if file_size == 0:
            errors.append("File is empty")
        
        return len(errors) == 0, errors
    
    def _extract_text_content(self, file_path: Path, file_extension: str) -> str:
        """Extract text content from file for indexing"""
        try:
            if file_extension == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif file_extension == ".pdf":
                return self._extract_pdf_text(file_path)
            elif file_extension == ".docx":
                return self._extract_docx_text(file_path)
            else:
                return ""
        except Exception as e:
            logger.warning(f"Failed to extract text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except ImportError:
            logger.warning("PyPDF2 not available for PDF text extraction")
            return ""
        except Exception as e:
            logger.warning(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            logger.warning("python-docx not available for DOCX text extraction")
            return ""
        except Exception as e:
            logger.warning(f"Error extracting DOCX text: {e}")
            return ""
    
    def create_document_request(self, db: Session, employee: Employee, 
                              request_data: Dict[str, Any]) -> Tuple[bool, Optional[DocumentRequest], List[str]]:
        """
        Create a new document request
        
        Args:
            db: Database session
            employee: Employee making the request
            request_data: Request data
            
        Returns:
            Tuple[bool, Optional[DocumentRequest], List[str]]: (success, request, errors)
        """
        try:
            request_id = f"DR{datetime.now().strftime('%Y%m%d')}{secrets.token_hex(3).upper()}"
            
            # Estimate completion time based on document type and urgency
            estimated_completion = self._estimate_completion_time(
                request_data.get("document_type"),
                request_data.get("urgency", "normal")
            )
            
            document_request = DocumentRequest(
                request_id=request_id,
                employee_id=employee.id,
                document_id=request_data.get("document_id"),
                document_title=request_data["document_title"],
                document_type=request_data["document_type"],
                description=request_data["description"],
                purpose=request_data.get("purpose"),
                format_preference=request_data.get("format_preference", "pdf"),
                delivery_method=request_data.get("delivery_method", "email"),
                urgency=request_data.get("urgency", "normal"),
                certified_copy=request_data.get("certified_copy", False),
                multiple_copies=request_data.get("multiple_copies", 1),
                special_instructions=request_data.get("special_instructions"),
                estimated_completion=estimated_completion,
                status=RequestStatus.PENDING
            )
            
            db.add(document_request)
            db.commit()
            db.refresh(document_request)
            
            logger.info(f"Document request created: {request_id}")
            return True, document_request, []
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error creating document request: {e}")
            return False, None, [f"Error creating request: {str(e)}"]
    
    def _estimate_completion_time(self, document_type: str, urgency: str) -> datetime:
        """Estimate completion time for document request"""
        base_hours = {
            "employment_certificate": 24,
            "salary_certificate": 24,
            "experience_letter": 48,
            "noc": 72,
            "other": 48
        }
        
        urgency_multiplier = {
            "urgent": 0.25,
            "high": 0.5,
            "normal": 1.0,
            "low": 2.0
        }
        
        hours = base_hours.get(document_type, 48) * urgency_multiplier.get(urgency, 1.0)
        return datetime.utcnow() + timedelta(hours=hours)
    
    def assign_request(self, db: Session, request: DocumentRequest, 
                      assignee: Employee) -> Tuple[bool, List[str]]:
        """Assign document request to HR personnel"""
        try:
            request.assigned_to = assignee.id
            request.status = RequestStatus.PROCESSING
            db.commit()
            
            logger.info(f"Request {request.request_id} assigned to {assignee.employee_id}")
            return True, []
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error assigning request: {e}")
            return False, [f"Error assigning request: {str(e)}"]
    
    def complete_request(self, db: Session, request: DocumentRequest,
                        completion_data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """Mark document request as completed"""
        try:
            request.status = RequestStatus.COMPLETED
            request.completed_at = datetime.utcnow()
            request.completion_notes = completion_data.get("notes")
            request.generated_file_path = completion_data.get("file_path")
            request.generated_file_name = completion_data.get("file_name")
            
            if completion_data.get("expiry_date"):
                request.expiry_date = completion_data["expiry_date"]
            
            db.commit()
            
            logger.info(f"Request {request.request_id} completed")
            return True, []
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error completing request: {e}")
            return False, [f"Error completing request: {str(e)}"]

# Global document service instance
document_service = DocumentService()